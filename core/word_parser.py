import re
from collections.abc import Iterator

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.inject import Inject


class WordParser:
    """Read an Exercise Director-style Word exercise pack."""

    INJECT_HEADING = re.compile(r"^No\.\s*(\d+)\s*$", re.IGNORECASE)
    CARD_FOOTER = re.compile(r"^Card\s+\d+\s+of\s+\d+", re.IGNORECASE)

    def __init__(self, filename):
        self.filename = filename
        self.document = None

    def open(self):
        self.document = Document(self.filename)

    def _require_document(self):
        if self.document is None:
            raise RuntimeError("Document has not been opened.")

    def get_summary(self):
        self._require_document()

        return {
            "paragraphs": len(self.document.paragraphs),
            "tables": len(self.document.tables),
            "injects": self.count_injects(),
        }

    def count_injects(self):
        return len(self.get_injects())

    def get_injects(self):
        """
        Build Inject objects from the MEL and the individual inject cards.
        """

        self._require_document()

        injects = self._read_mel()

        if not injects:
            return []

        injects_by_number = {
            inject.number: inject
            for inject in injects
        }

        self._read_inject_cards(injects_by_number)

        return injects

    def _read_mel(self):
        """
        Read number, time, phase, title and category from the MEL.
        The MEL is expected to be the first table in the document.
        """

        if not self.document.tables:
            return []

        mel_table = self.document.tables[0]
        injects = []

        for row in mel_table.rows[1:]:
            cells = row.cells

            if len(cells) < 5:
                continue

            number_text = self._clean_text(cells[0].text)

            if not number_text.isdigit():
                continue

            injects.append(
                Inject(
                    number=int(number_text),
                    exercise_time=self._clean_text(cells[1].text),
                    phase=self._clean_text(cells[2].text),
                    title=self._clean_text(cells[3].text),
                    category=self._clean_text(cells[4].text),
                )
            )

        return injects

    def _read_inject_cards(self, injects_by_number):
        """
        Walk through paragraphs and tables in their original document order.
        """

        blocks = list(self._iter_blocks(self.document))
        index = 0

        while index < len(blocks):
            block = blocks[index]

            if not isinstance(block, Table):
                index += 1
                continue

            inject_number = self._get_header_inject_number(block)

            if inject_number is None or inject_number not in injects_by_number:
                index += 1
                continue

            inject = injects_by_number[inject_number]

            card_blocks = []
            index += 1

            while index < len(blocks):
                next_block = blocks[index]

                if (
                    isinstance(next_block, Table)
                    and self._get_header_inject_number(next_block) is not None
                ):
                    break

                card_blocks.append(next_block)
                index += 1

            self._populate_inject_from_card(inject, card_blocks)

    def _populate_inject_from_card(self, inject, card_blocks):
        paragraphs = []

        for block in card_blocks:
            if isinstance(block, Table):
                self._read_metadata_table(inject, block)
            elif isinstance(block, Paragraph):
                text = self._clean_text(block.text)

                if text:
                    paragraphs.append(text)

        if not paragraphs:
            return

        # The first two paragraphs after the header are normally title and phase.
        if paragraphs:
            inject.title = paragraphs[0] or inject.title

        if len(paragraphs) > 1:
            inject.phase = paragraphs[1] or inject.phase

        sections = self._split_card_sections(paragraphs[2:])

        inject.inject_text = "\n\n".join(
            sections["inject_text"]
        ).strip()

        inject.expected_action = "\n\n".join(
            sections["expected_action"]
        ).strip()

        inject.facilitator_notes = "\n\n".join(
            sections["facilitator_notes"]
        ).strip()

        inject.attachments = sections["attachments"]

    def _read_metadata_table(self, inject, table):
        for row in table.rows:
            for cell in row.cells:
                text = self._clean_text(cell.text)

                if not text:
                    continue

                label, value = self._split_label_value(text)

                if label == "SOURCE":
                    inject.source = value
                elif label == "METHOD":
                    inject.method = value
                elif label == "AUDIENCE":
                    inject.audience = value
                elif label == "CATEGORY":
                    inject.category = value

    def _split_card_sections(self, paragraphs):
        sections = {
            "inject_text": [],
            "attachments": [],
            "expected_action": [],
            "facilitator_notes": [],
        }

        current_section = None

        for text in paragraphs:
            lowered = text.casefold()

            if lowered.startswith("inject text"):
                current_section = "inject_text"
                continue

            if lowered.startswith("attachment"):
                current_section = "attachments"
                continue

            if lowered.startswith("expected product / action"):
                current_section = "expected_action"
                continue

            if lowered.startswith("facilitators note"):
                current_section = "facilitator_notes"
                continue

            if self.CARD_FOOTER.match(text):
                current_section = None
                continue

            if current_section is None:
                continue

            if current_section == "attachments":
                sections["attachments"].extend(
                    self._extract_attachment_lines(text)
                )
            else:
                sections[current_section].append(text)

        return sections

    def _extract_attachment_lines(self, text):
        attachments = []

        for line in text.splitlines():
            cleaned = line.strip()

            if not cleaned:
                continue

            cleaned = cleaned.lstrip("☐").strip()

            if cleaned:
                attachments.append(cleaned)

        return attachments

    def _get_header_inject_number(self, table):
        if not table.rows or not table.rows[0].cells:
            return None

        first_cell = self._clean_text(table.rows[0].cells[0].text)
        match = self.INJECT_HEADING.match(first_cell)

        if not match:
            return None

        return int(match.group(1))

    @staticmethod
    def _split_label_value(text):
        match = re.match(
            r"^(SOURCE|METHOD|AUDIENCE|CATEGORY)\s*(.*)$",
            text,
            re.IGNORECASE | re.DOTALL,
        )

        if not match:
            return "", ""

        return match.group(1).upper(), match.group(2).strip()

    @staticmethod
    def _clean_text(text):
        lines = [
            line.strip()
            for line in text.replace("\xa0", " ").splitlines()
            if line.strip()
        ]

        return "\n".join(lines)

    @staticmethod
    def _iter_blocks(
        document: DocumentObject,
    ) -> Iterator[Paragraph | Table]:
        """
        Yield paragraphs and tables in their original Word document order.
        """

        body = document.element.body

        for child in body.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, document)
            elif isinstance(child, CT_Tbl):
                yield Table(child, document)